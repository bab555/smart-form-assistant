"""
Fast Tools - 程序直接执行的工具

特点：
1. 不经过 LLM，直接程序执行
2. 速度优先，毫秒级响应
3. 用于文件解析、数据填充、快速查询等

使用场景：
- Excel/CSV 解析
- Word 表格提取
- 图片 OCR
- 知识库快速查询
- 表格填充
"""
import pandas as pd
from io import BytesIO
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path

from app.core.logger import app_logger as logger
from app.services.product_index import CalibrationResult


@dataclass
class ParseResult:
    """解析结果"""
    success: bool
    rows: List[Dict[str, Any]]
    schema: List[Dict[str, str]]  # [{key, title, type}]
    message: str
    source_type: str  # excel/csv/word/image


class FastTools:
    """
    快速工具集
    
    所有方法都是同步/快速执行，不依赖 LLM
    """
    
    # ========== 文件解析工具 ==========
    
    @staticmethod
    def parse_excel(
        file_bytes: bytes,
        file_name: str = "file.xlsx",
        sheet_name: int = 0,
        max_rows: int = 1000
    ) -> ParseResult:
        """
        解析 Excel/CSV 文件
        
        Args:
            file_bytes: 文件字节
            file_name: 文件名（用于判断格式）
            sheet_name: 工作表索引
            max_rows: 最大行数限制
            
        Returns:
            ParseResult
        """
        try:
            ext = file_name.lower().split('.')[-1] if '.' in file_name else ''
            
            if ext == 'csv':
                df = pd.read_csv(BytesIO(file_bytes), nrows=max_rows)
                source_type = 'csv'
            else:
                df = pd.read_excel(BytesIO(file_bytes), sheet_name=sheet_name, nrows=max_rows)
                source_type = 'excel'
            
            # 清理列名
            df.columns = [str(col).strip() for col in df.columns]
            
            # 转换为行列表
            rows = df.to_dict('records')
            
            # 推断 schema
            schema = []
            if rows:
                first_row = rows[0]
                for key in first_row.keys():
                    value = first_row[key]
                    if isinstance(value, (int, float)) and not pd.isna(value):
                        col_type = "number"
                    else:
                        col_type = "text"
                    schema.append({
                        "key": key,
                        "title": key,
                        "type": col_type
                    })
            
            logger.info(f"[FastTools] Excel 解析: {len(rows)} 行, {len(schema)} 列")
            
            return ParseResult(
                success=True,
                rows=rows,
                schema=schema,
                message=f"成功解析 {len(rows)} 行数据",
                source_type=source_type
            )
            
        except Exception as e:
            logger.error(f"[FastTools] Excel 解析失败: {str(e)}")
            return ParseResult(
                success=False,
                rows=[],
                schema=[],
                message=f"解析失败: {str(e)}",
                source_type='excel'
            )
    
    @staticmethod
    def parse_word(file_bytes: bytes) -> ParseResult:
        """
        解析 Word 文档中的表格
        
        Args:
            file_bytes: 文件字节
            
        Returns:
            ParseResult
        """
        try:
            from docx import Document
            
            doc = Document(BytesIO(file_bytes))
            rows = []
            schema = []
            
            # 提取所有表格
            for table_idx, table in enumerate(doc.tables):
                # 第一行作为表头
                if table.rows:
                    header_row = table.rows[0]
                    headers = [cell.text.strip() or f"col_{i}" for i, cell in enumerate(header_row.cells)]
                    
                    if not schema:
                        schema = [{"key": h, "title": h, "type": "text"} for h in headers]
                    
                    # 后续行作为数据
                    for row in table.rows[1:]:
                        row_data = {}
                        for idx, cell in enumerate(row.cells):
                            key = headers[idx] if idx < len(headers) else f"col_{idx}"
                            row_data[key] = cell.text.strip()
                        if any(row_data.values()):
                            rows.append(row_data)
            
            if rows:
                logger.info(f"[FastTools] Word 解析: {len(rows)} 行表格数据")
                return ParseResult(
                    success=True,
                    rows=rows,
                    schema=schema,
                    message=f"成功提取 {len(rows)} 行表格数据",
                    source_type='word'
                )
            else:
                # 没有表格，提取段落文本
                text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                return ParseResult(
                    success=True,
                    rows=[],
                    schema=[],
                    message=text[:500] if text else "文档为空",
                    source_type='word_text'
                )
                
        except Exception as e:
            logger.error(f"[FastTools] Word 解析失败: {str(e)}")
            return ParseResult(
                success=False,
                rows=[],
                schema=[],
                message=f"解析失败: {str(e)}",
                source_type='word'
            )
    
    # ========== 知识库快速查询工具 ==========
    
    @staticmethod
    def quick_product_lookup(
        query: str,
        category: str = None,
        limit: int = 5
    ) -> List[Dict]:
        """
        快速商品查询（不经过 LLM）
        
        Args:
            query: 查询文本
            category: 分类筛选
            limit: 返回数量
            
        Returns:
            商品列表
        """
        from app.services.product_index import product_index
        
        results = product_index.search(query, limit=limit, category=category)
        
        return [
            {
                "id": r.product.id,
                "name": r.product.name,
                "code": r.product.code,
                "category": r.product.category,
                "unit": r.product.unit,
                "price": r.product.price,
                "score": r.score,
                "match_type": r.match_type,
            }
            for r in results
        ]
    
    @staticmethod
    def quick_calibrate(
        text: str,
        category: str = None
    ) -> CalibrationResult:
        """
        快速校准（不经过 LLM）
        
        Args:
            text: 待校准文本
            category: 分类提示
            
        Returns:
            CalibrationResult
        """
        from app.services.product_index import product_index
        return product_index.calibrate(text, category)
    
    @staticmethod
    def batch_calibrate(
        texts: List[str],
        category: str = None
    ) -> List[CalibrationResult]:
        """
        批量快速校准
        """
        from app.services.product_index import product_index
        return product_index.batch_calibrate(texts, category)
    
    # ========== 数据处理工具 ==========
    
    @staticmethod
    def extract_product_fields(row: Dict[str, Any]) -> Dict[str, Any]:
        """
        从行数据中提取商品相关字段
        
        自动识别常见字段名的变体
        """
        field_mappings = {
            "product": ["商品名", "商品名称", "品名", "产品名", "product", "name", "商品"],
            "quantity": ["数量", "数目", "qty", "quantity", "amount"],
            "unit": ["单位", "unit"],
            "price": ["单价", "价格", "price", "unit_price"],
            "total": ["金额", "总价", "total", "amount", "小计"],
        }
        
        result = {}
        row_lower = {k.lower(): v for k, v in row.items()}
        
        for target_key, candidates in field_mappings.items():
            for candidate in candidates:
                candidate_lower = candidate.lower()
                # 精确匹配
                if candidate_lower in row_lower:
                    result[target_key] = row_lower[candidate_lower]
                    break
                # 包含匹配
                for key in row_lower:
                    if candidate_lower in key:
                        result[target_key] = row_lower[key]
                        break
        
        return result
    
    @staticmethod
    def infer_schema(rows: List[Dict[str, Any]]) -> List[Dict[str, str]]:
        """
        从数据推断 schema
        """
        if not rows:
            return []
        
        first_row = rows[0]
        schema = []
        
        for key in first_row.keys():
            # 统计这一列的类型
            is_number = True
            for row in rows[:10]:  # 取前10行判断
                value = row.get(key)
                if value is not None and not isinstance(value, (int, float)):
                    try:
                        float(value)
                    except (ValueError, TypeError):
                        is_number = False
                        break
            
            schema.append({
                "key": key,
                "title": key,
                "type": "number" if is_number else "text"
            })
        
        return schema
    
    # ========== 计算工具 ==========
    
    @staticmethod
    def calculate_total(
        rows: List[Dict[str, Any]],
        price_field: str = None,
        quantity_field: str = None,
        total_field: str = None
    ) -> Dict[str, Any]:
        """
        计算表格统计数据
        
        Returns:
            {
                "row_count": 行数,
                "total_amount": 总金额,
                "total_quantity": 总数量,
                "breakdown": [每行小计]
            }
        """
        # 自动检测字段
        if not price_field or not quantity_field:
            field_candidates = {
                "price": ["单价", "价格", "price", "unit_price"],
                "quantity": ["数量", "数目", "qty", "quantity"],
                "total": ["金额", "总价", "total", "小计"],
            }
            
            if rows:
                keys = [k.lower() for k in rows[0].keys()]
                for field_type, candidates in field_candidates.items():
                    for c in candidates:
                        if c.lower() in keys:
                            if field_type == "price" and not price_field:
                                price_field = c
                            elif field_type == "quantity" and not quantity_field:
                                quantity_field = c
                            elif field_type == "total" and not total_field:
                                total_field = c
                            break
        
        total_amount = 0
        total_quantity = 0
        breakdown = []
        
        for row in rows:
            # 尝试获取小计
            row_total = 0
            if total_field and total_field in row:
                try:
                    row_total = float(row[total_field])
                except (ValueError, TypeError):
                    pass
            
            # 或者计算 price * quantity
            if row_total == 0 and price_field and quantity_field:
                try:
                    price = float(row.get(price_field, 0) or 0)
                    qty = float(row.get(quantity_field, 0) or 0)
                    row_total = price * qty
                except (ValueError, TypeError):
                    pass
            
            # 累计数量
            if quantity_field:
                try:
                    total_quantity += float(row.get(quantity_field, 0) or 0)
                except (ValueError, TypeError):
                    pass
            
            total_amount += row_total
            breakdown.append(row_total)
        
        return {
            "row_count": len(rows),
            "total_amount": round(total_amount, 2),
            "total_quantity": round(total_quantity, 2),
            "breakdown": breakdown,
        }


# 全局实例
fast_tools = FastTools()

